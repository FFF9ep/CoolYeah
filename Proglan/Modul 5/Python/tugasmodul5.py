from docx import Document
import os


def save_to_docx(pesanan_list):
    doc = Document()
    doc.add_heading('Manajemen Pesanan Restoran', 0)

    for pesanan in pesanan_list:
        doc.add_paragraph(f"ID Pesanan: {pesanan['id']}")
        doc.add_paragraph(f"Nama Pelanggan: {pesanan['nama']}")
        doc.add_paragraph(f"Menu Pesanan: {pesanan['menu']}")
        doc.add_paragraph(f"Jumlah: {pesanan['jumlah']}")
        doc.add_paragraph(f"Status: {pesanan['status']}")
        doc.add_paragraph(f"Image Path: {pesanan['image']}")
        doc.add_paragraph("-" * 40)

    doc.save("pesanan_restoran.docx")


def create_pesanan(pesanan_list):
    id_pesanan = len(pesanan_list) + 1
    nama = input("Nama Pelanggan: ")
    menu = input("Menu Pesanan: ")
    jumlah = int(input("Jumlah: "))
    image_path = input("Path Image: ")

    pesanan = {
        "id": id_pesanan,
        "nama": nama,
        "menu": menu,
        "jumlah": jumlah,
        "status": "diproses",
        "image": image_path
    }

    pesanan_list.append(pesanan)
    save_to_docx(pesanan_list)
    print("Pesanan berhasil dibuat!")


def read_pesanan(pesanan_list):
    if len(pesanan_list) == 0:
        print("Tidak ada pesanan.")
    else:
        for pesanan in pesanan_list:
            print(f"ID Pesanan: {pesanan['id']}")
            print(f"Nama Pelanggan: {pesanan['nama']}")
            print(f"Menu Pesanan: {pesanan['menu']}")
            print(f"Jumlah: {pesanan['jumlah']}")
            print(f"Status: {pesanan['status']}")
            print(f"Image Path: {pesanan['image']}")
            print("-" * 40)


def update_pesanan(pesanan_list):
    id_pesanan = int(input("Masukkan ID Pesanan yang ingin diubah: "))
    for pesanan in pesanan_list:
        if pesanan['id'] == id_pesanan:
            print(f"Pesanan {pesanan['id']} ditemukan.")
            new_status = input(
                "Masukkan status baru (diproses/selesai/batal): ")
            pesanan['status'] = new_status
            save_to_docx(pesanan_list)
            print(f"Status pesanan {
                  id_pesanan} telah diubah menjadi {new_status}.")
            return
    print("Pesanan tidak ditemukan.")


def delete_pesanan(pesanan_list):
    id_pesanan = int(input("Masukkan ID Pesanan yang ingin dihapus: "))
    for pesanan in pesanan_list:
        if pesanan['id'] == id_pesanan:
            if pesanan['status'] == 'batal':
                pesanan_list.remove(pesanan)
                save_to_docx(pesanan_list)
                print(f"Pesanan {id_pesanan} telah dihapus.")
                return
            else:
                print("Pesanan hanya bisa dihapus jika statusnya 'batal'.")
                return
    print("Pesanan tidak ditemukan.")


def search_pesanan(pesanan_list):
    nama_cari = input("Masukkan nama pelanggan yang ingin dicari: ")
    ditemukan = False
    for pesanan in pesanan_list:
        if nama_cari.lower() in pesanan['nama'].lower():
            print(f"ID Pesanan: {pesanan['id']}")
            print(f"Nama Pelanggan: {pesanan['nama']}")
            print(f"Menu Pesanan: {pesanan['menu']}")
            print(f"Jumlah: {pesanan['jumlah']}")
            print(f"Status: {pesanan['status']}")
            print(f"Image Path: {pesanan['image']}")
            print("-" * 40)
            ditemukan = True
    if not ditemukan:
        print("Pesanan tidak ditemukan.")


def main():
    pesanan_list = []

    if os.path.exists("pesanan_restoran.docx"):
        doc = Document("pesanan_restoran.docx")
        for para in doc.paragraphs:
            print(para.text)

    while True:
        print("\n1. Buat Pesanan")
        print("2. Baca Pesanan")
        print("3. Update Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan")
        print("6. Keluar")

        pilihan = input("Pilih menu (1/2/3/4/5/6): ")

        if pilihan == '1':
            create_pesanan(pesanan_list)
        elif pilihan == '2':
            read_pesanan(pesanan_list)
        elif pilihan == '3':
            update_pesanan(pesanan_list)
        elif pilihan == '4':
            delete_pesanan(pesanan_list)
        elif pilihan == '5':
            search_pesanan(pesanan_list)
        elif pilihan == '6':
            print("Terima kasih!")
            break
        else:
            print("Pilihan tidak valid, coba lagi.")


if __name__ == "__main__":
    main()
