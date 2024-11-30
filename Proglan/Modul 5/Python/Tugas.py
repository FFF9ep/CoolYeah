from docx import Document


def tambah_pesanan(nama, menu, jumlah):
    """Menambahkan pesanan baru ke dalam file."""
    try:
        document = Document("pesanan.docx")
    except:
        document = Document()
    document.add_paragraph(f"Nama: {nama}")
    document.add_paragraph(f"Menu: {menu}")
    document.add_paragraph(f"Jumlah: {jumlah}")
    document.add_paragraph("Status: Diproses")
    document.save("pesanan.docx")
    print("Pesanan berhasil ditambahkan!")


def tampilkan_pesanan():
    """Menampilkan semua pesanan yang ada dalam file."""
    try:
        document = Document("pesanan.docx")
        print("Data Pesanan:")
        for paragraph in document.paragraphs:
            print(paragraph.text)
    except:
        print("File pesanan tidak ditemukan.")


def update_pesanan(nama, status_baru):
    """Mengubah status pesanan."""
    try:
        document = Document("pesanan.docx")
        for paragraph in document.paragraphs:
            if paragraph.text.startswith(f"Nama: {nama}"):
                # Cari paragraf berikutnya dengan "Status:"
                for i in range(document.paragraphs.index(paragraph), len(document.paragraphs)):
                    if document.paragraphs[i].text.startswith("Status:"):
                        document.paragraphs[i].text = f"Status: {status_baru}"
                        document.save("pesanan.docx")
                        print("Status pesanan berhasil diubah!")
                        return
        print("Pesanan tidak ditemukan.")
    except:
        print("File pesanan tidak ditemukan.")


def hapus_pesanan(nama):
    """Menghapus pesanan berdasarkan nama."""
    try:
        document = Document("pesanan.docx")
        for paragraph in document.paragraphs:
            if paragraph.text.startswith(f"Nama: {nama}"):
                # Cari paragraf berikutnya dengan "Status:"
                for i in range(document.paragraphs.index(paragraph), len(document.paragraphs)):
                    if document.paragraphs[i].text.startswith("Status: Batal"):
                        # Hapus semua paragraf terkait pesanan ini
                        for j in range(4):
                            document.paragraphs.remove(
                                document.paragraphs[document.paragraphs.index(paragraph)])
                        document.save("pesanan.docx")
                        print("Pesanan berhasil dihapus!")
                        return
        print("Pesanan tidak ditemukan atau status tidak batal.")
    except:
        print("File pesanan tidak ditemukan.")


while True:
    print("\nManajemen Pesanan Restoran")
    print("1. Tambah Pesanan")
    print("2. Tampilkan Pesanan")
    print("3. Update Pesanan")
    print("4. Hapus Pesanan")
    print("5. Keluar")

    pilihan = input("Pilih menu: ")

    if pilihan == "1":
        nama = input("Masukkan Nama Pelanggan: ")
        menu = input("Masukkan Menu yang Dipesan: ")
        jumlah = input("Masukkan Jumlah Pesanan: ")
        tambah_pesanan(nama, menu, jumlah)
    elif pilihan == "2":
        tampilkan_pesanan()
    elif pilihan == "3":
        nama = input("Masukkan nama pelanggan yang ingin diupdate: ")
        status_baru = input("Masukkan status baru (Proses/Selesai/Batal): ")
        update_pesanan(nama, status_baru)
    elif pilihan == "4":
        nama = input("Masukkan nama pelanggan yang ingin dihapus: ")
        hapus_pesanan(nama)
    elif pilihan == "5":
        print("Terima kasih!")
        break
    else:
        print("Pilihan tidak valid.")
